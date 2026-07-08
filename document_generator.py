import os
import docx
from datetime import datetime
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Directory for saved transcriptions
OUTPUT_DIR = "transcriptions"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

def get_unicode_font():
    """Finds and registers a Unicode-capable system font on Windows."""
    system_fonts = [
        ("SegoeUI", "C:/Windows/Fonts/segoeui.ttf"),
        ("Arial", "C:/Windows/Fonts/arial.ttf"),
        ("Calibri", "C:/Windows/Fonts/calibri.ttf"),
    ]
    
    for name, path in system_fonts:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                return name
            except Exception:
                continue
    return "Helvetica" # Default fallback

FONT_NAME = get_unicode_font()

def save_to_docx(session_id, full_text_segments, summary_points=None, output_dir=None):
    """
    Saves/Updates the Word document with the full transcription.
    """
    save_dir = output_dir or OUTPUT_DIR
    file_path = os.path.join(save_dir, f"{session_id}.docx")
    
    doc = docx.Document()
    
    # Title
    title = doc.add_heading(level=0)
    title_run = title.add_run('Voice Transcription Report')
    title_run.font.name = 'Segoe UI' if FONT_NAME == 'SegoeUI' else 'Arial'
    
    # Metadata
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Session ID: {session_id}\n").bold = True
    p_meta.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    # Summary Section (if available)
    if summary_points:
        doc.add_heading('Executive Summary', level=1)
        for point in summary_points:
            doc.add_paragraph(point, style='List Bullet')
            
    # Full Transcription Section
    doc.add_heading('Full Transcription', level=1)
    
    if not full_text_segments:
        doc.add_paragraph("[No speech transcribed yet]")
    else:
        for seg in full_text_segments:
            # We assume seg is a dictionary: {"timestamp": "00:01 - 00:05", "text": "Hello"}
            # or a simple string.
            if isinstance(seg, dict):
                p = doc.add_paragraph()
                p.add_run(f"[{seg.get('timestamp', '')}] ").bold = True
                p.add_run(seg.get('text', ''))
            else:
                doc.add_paragraph(str(seg))
                
    doc.save(file_path)
    return file_path

def save_to_pdf(session_id, summary_points, full_text_preview="", output_dir=None):
    """
    Saves the summary and key points in a PDF document.
    """
    save_dir = output_dir or OUTPUT_DIR
    file_path = os.path.join(save_dir, f"{session_id}_summary.pdf")
    
    doc = SimpleDocTemplate(
        file_path,
        pagesize=letter,
        rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40
    )
    
    styles = getSampleStyleSheet()
    
    # Define custom styles using the registered Unicode font
    title_style = ParagraphStyle(
        'PDFTitle',
        parent=styles['Heading1'],
        fontName=FONT_NAME,
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#1e293b'),
        spaceAfter=15
    )
    
    section_style = ParagraphStyle(
        'PDFSection',
        parent=styles['Heading2'],
        fontName=FONT_NAME,
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#0f766e'),
        spaceBefore=15,
        spaceAfter=10
    )
    
    body_style = ParagraphStyle(
        'PDFBody',
        parent=styles['Normal'],
        fontName=FONT_NAME,
        fontSize=11,
        leading=16,
        textColor=colors.HexColor('#334155')
    )
    
    bullet_style = ParagraphStyle(
        'PDFBullet',
        parent=styles['Normal'],
        fontName=FONT_NAME,
        fontSize=11,
        leading=16,
        leftIndent=20,
        firstLineIndent=-10,
        textColor=colors.HexColor('#1e293b'),
        spaceAfter=5
    )
    
    story = []
    
    # Title
    story.append(Paragraph("Voice Transcription Summary", title_style))
    story.append(Spacer(1, 10))
    
    # Metadata Table
    meta_data = [
        [Paragraph("<b>Session ID:</b>", body_style), Paragraph(session_id, body_style)],
        [Paragraph("<b>Date:</b>", body_style), Paragraph(datetime.now().strftime('%Y-%m-%d %H:%M:%S'), body_style)]
    ]
    meta_table = Table(meta_data, colWidths=[100, 400])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
        ('PADDING', (0,0), (-1,-1), 8),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 20))
    
    # Summary Section
    story.append(Paragraph("Key Summary Points", section_style))
    if summary_points:
        for point in summary_points:
            story.append(Paragraph(f"• {point}", bullet_style))
    else:
        story.append(Paragraph("No summary points generated yet.", body_style))
        
    story.append(Spacer(1, 20))
    
    # Full Transcription Preview Section
    if full_text_preview:
        story.append(Paragraph("Transcription Preview", section_style))
        # Truncate preview if it's too long for summary PDF
        preview_text = full_text_preview
        if len(preview_text) > 800:
            preview_text = preview_text[:800] + " ... (Refer to Word document for full transcription)"
        story.append(Paragraph(preview_text, body_style))
        
    doc.build(story)
    return file_path
